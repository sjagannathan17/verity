"""Generate the two-page APA report (report/Verity_Report.docx).

Topic 2: Clinical Decision Making and Pattern Recognition. The report covers the
full topic and positions the Verity POC as the agentic decision-making / reasoning
layer, recommending it be paired with statistical pattern-recognition detection.

Run:  ../.venv-tools/bin/python tools/build_report.py   (from the verity/ dir)
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "report" / "Verity_Report.docx"

TITLE = "Agentic, Auditable Decision-Making for Payment Integrity"
SUBTITLE = (
    "Pairing Pattern Recognition with Chain-Reasoning AI across Treatment, Payment & Operations"
)
TAG = "Cotiviti GenAI Intern Assessment, Topic 2: Clinical Decision Making & Pattern Recognition"
AUTHOR = "Srinidhi Jagannathan   |   MS Business Analytics, Santa Clara University   |   June 30, 2026"

BODY = [
    ("Concept", [
        "Clinical decision making and pattern recognition span the Treatment, Payment, and "
        "Operations (TPO) lifecycle, and they divide naturally into two complementary halves. The "
        "first is statistical pattern recognition: classification, prediction, inference, "
        "clustering, and time-series anomaly detection, which surface WHAT looks abnormal in a "
        "stream of clinical or financial events. The second is decision making, namely chain "
        "reasoning and agentic generative AI, which determines WHAT TO DO about a flagged event and "
        "WHY, and produces a defensible record of that judgment. Payment integrity is the canonical "
        "TPO application: detectors surface suspicious claims, but a payer must still decide to pay, "
        "review, deny, or audit each one, and must justify that decision to providers and regulators.",
        "This analysis is anchored by a working proof of concept, Verity, that demonstrates the "
        "decision-making half of the topic. Given a submitted claim, an LLM orchestrator runs a chain "
        "of reasoning steps, invokes deterministic policy-check tools, and classifies the claim into a "
        "PAY/FLAG/DENY disposition with a cited rationale and a full audit trail. Crucially, the rules "
        "decide facts and the language model reasons and explains. This neuro-symbolic split keeps "
        "the decision auditable rather than opaque.",
    ]),
    ("Relevant Trends", [
        "Fraud, waste, and abuse (FWA) is a persistent, large-scale drain on U.S. health care "
        "spending (NHCAA, n.d.). On the pattern-recognition side, machine learning for FWA detection "
        "has matured from research into operations: supervised classification of claims, unsupervised "
        "clustering of providers with similar billing fingerprints, and anomaly detection of aberrant "
        "billing patterns are now standard tools (Bauder & Khoshgoftaar, 2017; Liu et al., 2008). Two "
        "of these techniques are worth naming concretely: prediction estimates the likelihood that a "
        "claim will be denied or that a provider is overbilling, in order to prioritize review, and "
        "inference is the model's per-claim conclusion drawn from the assembled rule evidence. These "
        "methods are strong at flagging that something is anomalous but weak at explaining why or "
        "deciding what should happen next.",
        "On the decision-making side, the frontier has shifted from black-box scoring toward "
        "explainable, agentic reasoning. Chain-of-thought prompting (Wei et al., 2022) and "
        "reason-and-act agent loops (Yao et al., 2023) let a model reason step by step over evidence "
        "and call external tools, while a deterministic rule layer supplies the facts. In regulated "
        "payment contexts the decisive trend is auditability and human-in-the-loop control: a "
        "recommendation is only useful if it cites the specific policy it applied and leaves a "
        "traceable record for appeal.",
    ]),
    ("Opportunities for Cotiviti", [
        "The opportunity is to connect the two halves: feed the output of Cotiviti's statistical "
        "detectors (pattern recognition, clustering, time-series anomaly detection) into an agentic "
        "reasoning layer that classifies each candidate into a disposition, explains it in plain "
        "language, and cites the governing rule. This raises analyst throughput by pre-screening and "
        "drafting rationales, improves consistency because every decision carries cited evidence, and "
        "strengthens appeals defensibility because the audit trail is the record. The Verity POC shows "
        "this concretely, including a 'near-miss save': a same-day repeat procedure that trips a "
        "units edit but is cleared because a modifier documents a legitimate repeat, demonstrating "
        "false-positive avoidance and provider-abrasion sensitivity, the maturity that separates a "
        "decision system from a blunt auto-denier.",
    ]),
    ("Threats and Risks", [
        "The dominant risk is letting a generative model decide facts: LLMs can hallucinate or "
        "over-confidently justify an incorrect denial, which in payment creates compliance exposure "
        "and member harm. Verity mitigates this structurally: deterministic rules decide, and the "
        "model only reasons and explains. Statistical detection introduces its own risk: false "
        "positives drive provider abrasion, so the reasoning-and-review layer is not optional. "
        "Protected health information demands strict governance (encryption, access control, "
        "minimized prompts). Finally, both models and policies drift, as detector accuracy decays "
        "and coding rules change, so an evaluation pipeline (precision/recall per disposition on a "
        "labeled gold set) and ongoing monitoring are prerequisites, not afterthoughts. Fidelity of "
        "the encoded policy also matters: edit semantics, for example line-level versus "
        "date-of-service unit limits (CMS, n.d.), must be exact, because a wrong conversion yields "
        "wrong denials at scale.",
    ]),
    ("Strategic Recommendation", [
        "Cotiviti should invest in a two-layer payment-integrity architecture that explicitly unites "
        "the topic's techniques. Layer one is statistical pattern recognition (classification, "
        "clustering, and time-series anomaly detection) to surface candidate claims and providers. "
        "Layer two is an agentic, rule-grounded reasoning engine that classifies each candidate into a "
        "disposition, produces a cited rationale, and routes denials and audits through human sign-off. "
        "Verity is a working demonstration of layer two end to end: chain reasoning, agentic "
        "tool-calling, claim classification, and an exportable audit trail, with an offline-safe mode "
        "for reliability.",
        "The recommended path is deliberately staged: begin with a bounded, high-confidence edit set "
        "(as the POC does), measure precision and recall against a labeled gold set, and scale only as "
        "the metrics justify. This keeps the system auditable and conservative while delivering "
        "immediate value, and it positions the reasoning layer to consume whatever detection signals "
        "Cotiviti already operates rather than replacing them, extending from payment accuracy into "
        "adjacent surfaces such as risk adjustment and quality and Star ratings.",
    ]),
]

REFERENCES = [
    "Bauder, R. A., & Khoshgoftaar, T. M. (2017). Medicare fraud detection using machine learning "
    "methods. In 2017 16th IEEE International Conference on Machine Learning and Applications (ICMLA) "
    "(pp. 858-865). IEEE. https://doi.org/10.1109/ICMLA.2017.00-48",
    "Centers for Medicare & Medicaid Services. (n.d.). National Correct Coding Initiative (NCCI) "
    "policy manual and Medically Unlikely Edits (MUE). U.S. Department of Health and Human Services. "
    "https://www.cms.gov/medicare/coding-billing/national-correct-coding-initiative-ncci-edits",
    "Liu, F. T., Ting, K. M., & Zhou, Z.-H. (2008). Isolation forest. In 2008 Eighth IEEE "
    "International Conference on Data Mining (pp. 413-422). IEEE. https://doi.org/10.1109/ICDM.2008.17",
    "National Health Care Anti-Fraud Association. (n.d.). The challenge of health care fraud. "
    "Retrieved from https://www.nhcaa.org",
    "Wei, J., Wang, X., Schuurmans, D., Bosma, M., Ichter, B., Xia, F., Chi, E., Le, Q., & Zhou, D. "
    "(2022). Chain-of-thought prompting elicits reasoning in large language models. Advances in "
    "Neural Information Processing Systems, 35, 24824-24837.",
    "Yao, S., Zhao, J., Yu, D., Du, N., Shafran, I., Narasimhan, K., & Cao, Y. (2023). ReAct: "
    "Synergizing reasoning and acting in language models. In International Conference on Learning "
    "Representations (ICLR).",
]


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(72)
        section.left_margin = section.right_margin = Pt(72)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x5A, 0x2D, 0x8C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(SUBTITLE)
    rs.italic = True
    rs.font.size = Pt(11)

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = tag.add_run(TAG)
    rt.font.size = Pt(9.5)
    rt.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(AUTHOR)
    rm.font.size = Pt(9.5)
    rm.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for heading, paras in BODY:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(8)
        hr = h.add_run(heading)
        hr.bold = True
        hr.font.size = Pt(12.5)
        hr.font.color.rgb = RGBColor(0x5A, 0x2D, 0x8C)
        for p in paras:
            para = doc.add_paragraph(p)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()
    refh = doc.add_paragraph()
    refh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = refh.add_run("References")
    rr.bold = True
    rr.font.size = Pt(12.5)
    for ref in REFERENCES:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.first_line_indent = Pt(-36)
        p.paragraph_format.space_after = Pt(8)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    words = sum(len(p.split()) for _, ps in BODY for p in ps)
    print(f"Wrote {OUT} (~{words} body words)")


if __name__ == "__main__":
    main()
